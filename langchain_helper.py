import os
import fitz  
import docx
import pandas as pd

from dotenv import load_dotenv
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.chat_models import ChatOpenAI
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate

load_dotenv()

def extract_text(file_obj, filename):
    ext = filename.split(".")[-1].lower()

    if ext == "pdf":
        from PyPDF2 import PdfReader
        reader = PdfReader(file_obj)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text

    elif ext in ["docx", "doc"]:
        import docx
        doc = docx.Document(file_obj)
        return "\n".join([para.text for para in doc.paragraphs])

    elif ext == "txt":
        return file_obj.read().decode("utf-8")

    else:
        raise ValueError(f"Unsupported file type: {ext}")



def create_faiss_index(text, faiss_dir):
    """
    Create a FAISS index from extracted text and save to faiss_dir.
    """
    openai_key = os.getenv("OPENAI_API_KEY")
    embeddings = OpenAIEmbeddings(openai_api_key=openai_key)

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=150,
        separators=["\n\n", "\n", ".", " ", ""]
    )

    chunks = text_splitter.split_text(text)
    documents = [Document(page_content=chunk) for chunk in chunks]

    vectorstore = FAISS.from_documents(documents, embeddings)

    os.makedirs(faiss_dir, exist_ok=True)
    vectorstore.save_local(faiss_dir)

    print(f"✅ FAISS index saved to: {faiss_dir}")


def get_qa_chain(faiss_path):
    openai_key = os.getenv("OPENAI_API_KEY")
    embeddings = OpenAIEmbeddings(openai_api_key=openai_key)
    llm = ChatOpenAI(model_name="gpt-3.5-turbo", openai_api_key=openai_key)

    if not os.path.exists(faiss_path):
        raise ValueError(f"FAISS index not found at {faiss_path}")


    vectorstore = FAISS.load_local(
        faiss_path,
        embeddings,
        allow_dangerous_deserialization=True 
    )
    retriever = vectorstore.as_retriever(search_kwargs={"k": 5})

    prompt_template = PromptTemplate.from_template("""
    You are a document analyzer. Answer the question in detail based on the context.
    If the question is not related to the document, reply with "I don't know".

    Context:
    {context}

    Question: {question}
    Answer:
    """)

    qa_chain = RetrievalQA.from_chain_type(
        llm=llm,
        retriever=retriever,
        chain_type="stuff",
        chain_type_kwargs={"prompt": prompt_template},
        return_source_documents=True
    )

    return qa_chain
